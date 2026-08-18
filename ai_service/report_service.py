"""Generate evidence-backed Word reports from the existing dashboard APIs."""

from __future__ import annotations

import json
import tempfile
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from langchain_core.messages import HumanMessage, SystemMessage

from ai_service.agent import _build_model, _content_to_text
from ai_service.schemas import ReportRequest, ToolCall
from ai_service.tools import ToolRequestError, call_tool_with_filters


REPORT_TOOL_GROUPS: dict[str, list[str]] = {
    "comprehensive": [
        "get_kpi", "get_age_gender", "get_payment", "get_disposition", "get_medical_surgical",
        "get_admission_emergency", "get_hospital_ranking", "get_disease_systems",
        "get_top_diagnoses", "get_severity",
    ],
    "operations": ["get_kpi", "get_hospital_ranking"],
    "patient": ["get_kpi", "get_age_gender", "get_payment", "get_disposition", "get_admission_emergency", "get_medical_surgical"],
    "disease": ["get_kpi", "get_disease_systems", "get_top_diagnoses", "get_severity"],
}

REPORT_TITLES = {
    "comprehensive": "住院数据综合分析报告",
    "operations": "医院运营分析报告",
    "patient": "患者结构分析报告",
    "disease": "疾病与病情风险分析报告",
}

REPORT_DIR = Path(tempfile.gettempdir()) / "medical-agent-reports"
REPORT_DIR.mkdir(parents=True, exist_ok=True)


def _as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _as_list(value: Any) -> list[dict[str, Any]]:
    return [item for item in value if isinstance(item, dict)] if isinstance(value, list) else []


def _number(value: Any) -> str:
    if value is None:
        return "暂无"
    if isinstance(value, float):
        return f"{value:,.2f}"
    try:
        return f"{int(value):,}"
    except (TypeError, ValueError):
        return str(value)


def _percent(value: Any) -> str:
    if value is None:
        return "暂无"
    try:
        number = float(value)
        if number <= 1:
            number *= 100
        return f"{number:.2f}%"
    except (TypeError, ValueError):
        return str(value)


def _money(value: Any) -> str:
    if value is None:
        return "暂无"
    try:
        return f"${float(value):,.2f}"
    except (TypeError, ValueError):
        return str(value)


def _scope_text(request: ReportRequest) -> str:
    filters = request.filters
    parts = [filters.hospital_service_area, filters.hospital_county, filters.facility_name]
    labels = ["区域", "县", "医院"]
    selected = [f"{label}：{value}" for label, value in zip(labels, parts) if value]
    return "；".join(selected) if selected else "全量数据范围"


def _section(title: str, summary: str, findings: list[str], metrics: list[tuple[str, str]], source_tools: list[str], tables: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    return {
        "title": title,
        "summary": summary,
        "findings": findings,
        "metrics": [{"label": label, "value": value} for label, value in metrics],
        "source_tools": source_tools,
        "tables": tables or [],
    }


def _build_sections(report_type: str, results: dict[str, Any]) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    kpi = _as_dict(results.get("get_kpi"))
    if kpi:
        sections.append(_section(
            "一、核心运营概况",
            "当前筛选范围内的医院运营核心指标如下。",
            [
                f"当前共有 {_number(kpi.get('hospital_count'))} 家医院，出院记录 {_number(kpi.get('discharge_count'))} 条。",
                f"急诊有效记录占比为 {_percent(kpi.get('emergency_ratio'))}。",
            ],
            [("医院数", _number(kpi.get("hospital_count"))), ("出院量", _number(kpi.get("discharge_count"))),
             ("平均住院天数", _number(kpi.get("average_length_of_stay"))), ("总收费", _money(kpi.get("total_charges"))),
             ("总成本", _money(kpi.get("total_costs"))), ("急诊占比", _percent(kpi.get("emergency_ratio")))],
            ["get_kpi"],
        ))

    if report_type in {"comprehensive", "operations"}:
        ranking = _as_list(results.get("get_hospital_ranking"))
        top = ranking[0] if ranking else {}
        sections.append(_section(
            "二、医院运营表现",
            "医院排名按照当前分析接口返回的出院量排序。",
            [f"出院量最高的医院为 {top.get('facility_name', '暂无')}，出院量 {_number(top.get('discharge_count'))}。"],
            [("排名医院数", _number(len(ranking))), ("第一名出院量", _number(top.get("discharge_count"))),
             ("第一名平均成本", _money(top.get("average_cost")))],
            ["get_hospital_ranking"],
            [{
                "title": "医院运营 Top 10",
                "headers": ["医院名称", "所在县", "出院量", "平均成本"],
                "rows": [[str(item.get("facility_name", "暂无")), str(item.get("hospital_county", "暂无")), _number(item.get("discharge_count")), _money(item.get("average_cost"))] for item in ranking],
            }],
        ))

    if report_type in {"comprehensive", "patient"}:
        payments = _as_list(results.get("get_payment"))
        top_payment = max(payments, key=lambda item: float(item.get("discharge_count") or 0), default={})
        ages = _as_list(results.get("get_age_gender"))
        medical_surgical = _as_list(results.get("get_medical_surgical"))
        top_age = max(ages, key=lambda item: float(item.get("discharge_count") or 0), default={})
        sections.append(_section(
            "三、患者结构",
            "患者结构根据年龄性别、支付方式、入院急诊和离院去向汇总结果分析。",
            [f"出院量最多的年龄性别组合为 {top_age.get('age_group', '暂无')} / {top_age.get('gender', '暂无')}。",
             f"主要支付方式为 {top_payment.get('payment_typology_1', '暂无')}，对应出院量 {_number(top_payment.get('discharge_count'))}。"],
            [("年龄性别组合数", _number(len(ages))), ("支付方式数", _number(len(payments))),
             ("主要支付方式", str(top_payment.get("payment_typology_1", "暂无")))],
            ["get_age_gender", "get_payment", "get_disposition", "get_admission_emergency", "get_medical_surgical"],
            [{
                "title": "主要支付方式",
                "headers": ["支付方式", "出院量", "总收费", "总成本"],
                "rows": [[str(item.get("payment_typology_1", "暂无")), _number(item.get("discharge_count")), _money(item.get("total_charges")), _money(item.get("total_costs"))] for item in payments],
            }, {
                "title": "内外科结构",
                "headers": ["分类", "出院量", "总收费", "总成本"],
                "rows": [[str(item.get("apr_medical_surgical_description", "暂无")), _number(item.get("discharge_count")), _money(item.get("total_charges")), _money(item.get("total_costs"))] for item in medical_surgical],
            }],
        ))

    if report_type in {"comprehensive", "disease"}:
        systems = _as_list(results.get("get_disease_systems"))
        diagnoses = _as_list(results.get("get_top_diagnoses"))
        severity = _as_list(results.get("get_severity"))
        top_system = max(systems, key=lambda item: float(item.get("discharge_count") or 0), default={})
        top_diagnosis = max(diagnoses, key=lambda item: float(item.get("discharge_count") or 0), default={})
        sections.append(_section(
            "四、疾病与病情风险",
            "疾病部分按 APR MDC 疾病系统、CCSR 诊断和病情严重程度汇总。",
            [f"出院量最高的疾病系统为 {top_system.get('apr_mdc_code', '暂无')}，出院量 {_number(top_system.get('discharge_count'))}。",
             f"高发诊断 Top 10 中记录数最高的诊断为 {top_diagnosis.get('ccsr_diagnosis_code', '暂无')}。"],
            [("疾病系统数", _number(len(systems))), ("高发诊断数", _number(len(diagnoses))), ("严重程度组合数", _number(len(severity)))],
            ["get_disease_systems", "get_top_diagnoses", "get_severity"],
            [{
                "title": "高发疾病 Top 10",
                "headers": ["诊断编码", "诊断描述", "出院量"],
                "rows": [[str(item.get("ccsr_diagnosis_code", "暂无")), str(item.get("ccsr_diagnosis_description", "暂无")), _number(item.get("discharge_count"))] for item in diagnoses],
            }],
        ))
    return sections


async def _generate_executive_summary(request: ReportRequest, sections: list[dict[str, Any]]) -> str:
    fallback = f"本报告覆盖{_scope_text(request)}，基于平台汇总数据生成。报告共包含 {len(sections)} 个分析章节，结论仅用于数据观察和运营分析。"
    try:
        model = _build_model()
        prompt = (
            "你是医疗数据分析报告撰写助手。只能根据给出的结构化事实写中文执行摘要，不能虚构数字，"
            "不能给出诊疗建议。摘要控制在 180 字以内，说明分析范围、主要发现和数据用途。\n"
            f"分析范围：{_scope_text(request)}\n"
            f"用户关注重点：{request.focus or '无额外重点'}\n"
            f"结构化章节：{json.dumps(sections, ensure_ascii=False, default=str)}"
        )
        response = await model.ainvoke([SystemMessage(content="你只负责撰写数据分析摘要。"), HumanMessage(content=prompt)])
        return _content_to_text(response.content) or fallback
    except Exception:
        return fallback


async def collect_report(request: ReportRequest) -> tuple[dict[str, Any], list[ToolCall]]:
    """Collect approved endpoint data and construct a report model."""

    results: dict[str, Any] = {}
    calls: list[ToolCall] = []
    filters = request.filters.model_dump()
    for name in REPORT_TOOL_GROUPS[request.report_type]:
        try:
            results[name] = await call_tool_with_filters(name, filters)
            calls.append(ToolCall(name=name, status="success"))
        except Exception:
            calls.append(ToolCall(name=name, status="failed"))
    if not any(call.status == "success" for call in calls):
        raise ToolRequestError("报告所需的医疗分析数据暂时不可用。")
    sections = _build_sections(request.report_type, results)
    report = {
        "title": REPORT_TITLES[request.report_type],
        "report_type": request.report_type,
        "scope": request.filters.model_dump(),
        "generated_at": datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds"),
        "executive_summary": await _generate_executive_summary(request, sections),
        "sections": sections,
        "limitations": [
            "数据来源为去标识化住院出院记录。",
            "部分医院区域、县或字段可能为空，缺失值不代表零值。",
            "平均住院天数、平均收费和平均成本按当前筛选范围的汇总值除以出院记录数计算。",
            "本报告用于统计分析和运营观察，不构成临床诊断或治疗建议。",
        ],
    }
    return report, calls


def _set_cell(cell: Any, text: str, bold: bool = False) -> None:
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.bold = bold


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            _set_cell(cell, value)


def render_docx(report: dict[str, Any], tool_calls: list[ToolCall]) -> tuple[str, Path]:
    """Render a report to a temporary DOCX and return its short-lived identifier."""

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report["title"])
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"分析范围：{_scope_text(ReportRequest.model_validate({'filters': report['scope']}))}\n").bold = True
    meta.add_run(f"生成时间：{report['generated_at']}")

    document.add_heading("执行摘要", level=1)
    document.add_paragraph(report["executive_summary"])
    for section_data in report["sections"]:
        document.add_heading(section_data["title"], level=1)
        document.add_paragraph(section_data["summary"])
        document.add_paragraph(f"数据来源：{', '.join(section_data.get('source_tools', [])) or '无'}")
        metrics = section_data.get("metrics", [])
        if metrics:
            _add_table(document, ["指标", "数值"], [[item["label"], item["value"]] for item in metrics])
        for table_data in section_data.get("tables", []):
            document.add_paragraph(table_data["title"]).runs[0].bold = True
            _add_table(document, table_data["headers"], table_data["rows"])
        for finding in section_data.get("findings", []):
            document.add_paragraph(finding, style="List Bullet")
    document.add_heading("数据范围与使用限制", level=1)
    for item in report["limitations"]:
        document.add_paragraph(item, style="List Bullet")
    document.add_paragraph(f"数据工具调用：{', '.join(call.name for call in tool_calls if call.status == 'success') or '无'}")

    report_id = uuid.uuid4().hex
    path = REPORT_DIR / f"medical_report_{report_id}.docx"
    document.save(path)
    return report_id, path


async def generate_report_document(request: ReportRequest) -> tuple[str, Path, dict[str, Any], list[ToolCall]]:
    report, calls = await collect_report(request)
    report_id, path = render_docx(report, calls)
    return report_id, path, report, calls
